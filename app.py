import streamlit as st
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from datetime import datetime, timedelta
import io

# 📘 Configurações iniciais
st.set_page_config(page_title="Plano Quinzenal Inteligente", layout="wide")
st.title("📘 Gerador Inteligente de Plano Quinzenal com Lições e Disciplinas")

# 📖 Plano Analítico (exemplo)
plano_analitico = {
    "Matemática": {
        11: "Multiplicação de números naturais",
        12: "Divisão de números naturais",
        13: "Problemas envolvendo as quatro operações"
    },
    "Português": {
        5: "Leitura e interpretação de texto",
        6: "Uso de pontuação",
        7: "Produção de texto narrativo"
    },
    "Ciências": {
        1: "O corpo humano",
        2: "Sistema digestivo",
        3: "Sistema respiratório"
    }
}

# 📅 Configurações do plano
col1, col2 = st.columns(2)
with col1:
    data_inicio = st.date_input("Data de início da quinzena", value=datetime.today())
with col2:
    curriculo_local = st.text_input("Currículo local", value="Currículo de Moçambique")

# 📆 Dias letivos (máximo 20 para limitar a 4 páginas)
dias_letivos = [data_inicio + timedelta(days=i) for i in range(14) if (data_inicio + timedelta(days=i)).weekday() < 5]
max_dias = min(len(dias_letivos), 20)

# 🧠 Função para buscar conteúdos das lições
def buscar_conteudo(disciplina, licoes_str):
    if disciplina not in plano_analitico:
        return "Disciplina não encontrada no plano analítico."
    licoes = []
    for parte in licoes_str.replace(" ", "").split(","):
        if "-" in parte:
            ini, fim = parte.split("-")
            licoes.extend(range(int(ini), int(fim)+1))
        else:
            licoes.append(int(parte))
    conteudos = [plano_analitico[disciplina].get(num, f"Lição {num} não encontrada") for num in licoes]
    return "; ".join(conteudos)

# 📋 Interface por dia
st.subheader("📋 Preenchimento Diário")
plano_diario = []

for dia in dias_letivos[:max_dias]:
    st.markdown(f"### 📅 {dia.strftime('%A, %d/%m/%Y')}")
    tempos_dia = []
    for tempo in range(1, 7):
        col1, col2, col3 = st.columns([1, 2, 2])
        horario = col1.text_input(f"{dia.strftime('%d/%m')} - Tempo {tempo} - Horário", value=f"{7+tempo}:30 - {8+tempo}:15")
        disciplina = col2.selectbox(f"{dia.strftime('%d/%m')} - Tempo {tempo} - Disciplina", options=list(plano_analitico.keys()), key=f"{dia}_disc_{tempo}")
        licoes = col3.text_input(f"{dia.strftime('%d/%m')} - Tempo {tempo} - Lição(ões)", value="11", key=f"{dia}_lic_{tempo}")
        orientacao = buscar_conteudo(disciplina, licoes)
        tempos_dia.append({
            "Tempo": tempo,
            "Horário": horario,
            "Disciplina": disciplina,
            "Licoes": licoes,
            "Orientacao": orientacao
        })
    plano_diario.append({"dia": dia, "tempos": tempos_dia})

# 👀 Pré-visualização
st.subheader("👀 Prévia do Plano Quinzenal")
for dia in plano_diario:
    st.markdown(f"## 📅 {dia['dia'].strftime('%A, %d/%m/%Y')}")
    for tempo in dia["tempos"]:
        st.markdown(f"""
        **Tempo {tempo['Tempo']}**  
        🕒 Horário: {tempo['Horário']}  
        📘 Disciplina: {tempo['Disciplina']}  
        🔢 Lição(ões): {tempo['Licoes']}  
        📖 Conteúdo: {tempo['Orientacao']}
        """)

# 📄 Gerar documento Word
if st.button("📄 Baixar Plano Quinzenal"):
    doc = Document()

    # 🧭 Orientação horizontal
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    # 📝 Cabeçalho
    doc.add_heading("Plano Quinzenal Escolar", 0)
    doc.add_paragraph(f"Currículo Local: {curriculo_local}")
    doc.add_paragraph(f"Período: {data_inicio.strftime('%d/%m/%Y')} até {(data_inicio + timedelta(days=13)).strftime('%d/%m/%Y')}")
    doc.add_paragraph(" ")

    # 📅 Grelha por dia
    for dia in plano_diario:
        doc.add_heading(f"📅 {dia['dia'].strftime('%A, %d/%m/%Y')}", level=1)
        tabela = doc.add_table(rows=1, cols=5)
        tabela.style = 'Table Grid'
        hdr = tabela.rows[0].cells
        hdr[0].text = "Tempo"
        hdr[1].text = "Horário"
        hdr[2].text = "Disciplina"
        hdr[3].text = "Lição(ões)"
        hdr[4].text = "Conteúdo"

        for tempo in dia["tempos"]:
            linha = tabela.add_row().cells
            linha[0].text = str(tempo["Tempo"])
            linha[1].text = tempo["Horário"]
            linha[2].text = tempo["Disciplina"]
            linha[3].text = tempo["Licoes"]
            linha[4].text = tempo["Orientacao"]

        doc.add_paragraph(" ")

    # 📤 Exportar
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    st.download_button("📥 Baixar plano quinzenal", data=buffer, file_name="plano_quinzenal_inteligente.docx")
